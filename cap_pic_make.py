from docx import Document
from docx.shared import Inches


from docx import Document
import docx

#  table: docx.table.Table,
# def add_caption(doc, caption: str,target = 'Table'):
#     """
#     Based on: https://github.com/python-openxml/python-docx/issues/359
#     """
#     # target = 'Table'
#     # doc
#     # caption type
#     paragraph = doc.add_paragraph(f'{target} ', style='Caption')

#     # numbering field
#     run = paragraph.add_run()

#     fldChar = docx.oxml.OxmlElement('w:fldChar')
#     fldChar.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
#     run._r.append(fldChar)

#     instrText = docx.oxml.OxmlElement('w:instrText')
#     # instrText.text = f' SEQ {target} \\* ARABIC'
#     instrText.text = f' SEQ {target} \\* ARABIC'
#     # 表 0.2
#     run._r.append(instrText)

#     fldChar = docx.oxml.OxmlElement('w:fldChar')
#     fldChar.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
#     run._r.append(fldChar)

#     # caption text
#     paragraph.add_run(f' {caption}')

import word_util
document = Document()

# 插入图片
# document.add_picture('monty-truth.png', width=Inches(1.25))
document.add_paragraph('这是一张蒙提·派森的真相海报')
# 插入图片题注
# Caption
# Table  这是一个表格
p = document.add_paragraph()
p.add_run('图').bold = True
p.add_run(': 这是一张蒙提·派森的真相海报')
document.add_paragraph('这是一张蒙提·派森的真相海报')

# 插入表格
table = document.add_table(rows=2, cols=2)
cell = table.cell(0, 1)
cell.text = '表格内容'
cell = table.cell(1, 0)
cell.text = '表格内容'

# 插入表格题注
p = document.add_paragraph()
p.add_run('表').bold = True
p.add_run(': 这是一个表格')
document.add_paragraph('这是一个表格')
word_util.add_caption(document, '这是一个表格',target = '表')
import time_util
now_time_str=time_util.get_now_time_str()
out_file_name=fr'D:\bishe\doc_cap_demo\demo_{now_time_str}.docx'

# out_file_name='demo2.docx'
document.save(out_file_name)

import os 
os.system(f'start {out_file_name}')

# 您可以使用Python-docx库来读取Word文档中的所有题注。下面是一个示例代码，它演示了如何读取Word文档中的所有题注。12

document = Document(f'{out_file_name}')

for paragraph in document.paragraphs:
    # print(paragraph)
    print(paragraph.style.name)
    if paragraph.style.name.startswith('Caption'):
        print(paragraph.text)
# 希望这可以帮助到您！如果您有任何其他问题，请随时问我。

# 已收到消息. 您可以使用Python-docx库来读取Word文档中的所有题注。下面是一个示例代码，它演示了如何读取Word文档中的所有题注。 ```python from docx import Document document = Document('document.docx') for paragraph in document.paragraphs: if paragraph.style.name.startswith('Caption'): print(paragraph.text) ``` 希望这可以帮助到您！如果您有任何其他问题，请随时问我。
# 了解详细信息:
# 1. blog.csdn.net
# 2. blog.csdn.net
# 3. zhuanlan.zhihu.com