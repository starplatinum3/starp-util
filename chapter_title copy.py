import docx
import re

def read_docx(filename):
    # 创建Document对象
    doc = docx.Document(filename)

    # 初始化标题计数器和上一个标题等级
    chapter_num = 0
    last_level = 0
    section_num = 0
    last_section_num = 0

    # 遍历Document对象中的所有Paragraph对象
    for para in doc.paragraphs:
        # 获取当前段落的样式
        style = para.style.name

        # 如果样式是标题，解析标题内容并打印
        if 'Heading' in style:
            level = int(style[-1])
            text = para.text

            # 如果标题以数字开头，表示章节标题
            match = re.match(r'^(\d+)(.*?)$', text.strip())
            if match:
                chapter_num = int(match.group(1))
                section_num = 0
                last_section_num = 0
                print(f"{chapter_num}. {match.group(2)}")
                last_level = level
                continue

            # 如果标题以字母开头，表示小节标题
            match = re.match(r'^([A-Za-z]+)(.*?)$', text.strip())
            if match:
                section_num += 1
                if level > last_level:
                    section_num = last_section_num + 1
                last_section_num = section_num

                # 按照格式打印标题
                print(f"{chapter_num}.{section_num} {match.group(1)}{match.group(2)}")
                last_level = level
                continue


import docx
import re

import time_util
import os

def apply_heading_style(filename):
    # 创建Document对象
    doc = docx.Document(filename)

    # 初始化标题计数器和上一个标题等级
    chapter_num = 0
    last_level = 0
    section_num = 0
    last_section_num = 0

    # 遍历Document对象中的所有Paragraph对象
    for para in doc.paragraphs:
        # 获取当前段落的样式
        style = para.style.name

        # 如果样式是标题，解析标题内容并修改样式
        if 'Heading' in style:
            level = int(style[-1])
            text = para.text
            # 如果他是类似 第1章 绪论，我认为他是一级别的标题
            # 如果标题以数字开头，表示章节标题
            match = re.match(r'^(\d+)(.*?)$', text.strip())
            if match:
                chapter_num = int(match.group(1))
                section_num = 0
                last_section_num = 0
                para.style = doc.styles[f"Heading {level}"]
                para.text = f"{chapter_num}. {match.group(2)}"
                last_level = level
                continue

            # 如果标题以字母开头，表示小节标题
            match = re.match(r'^([A-Za-z]+)(.*?)$', text.strip())
            if match:
                section_num += 1
                if level > last_level:
                    section_num = last_section_num + 1
                last_section_num = section_num

                # 修改样式并设置内容
                para.style = doc.styles[f"Heading {level}"]
                para_text=f"{chapter_num}.{section_num} {match.group(1)}{match.group(2)}"
                print("para_text")
                print(para_text)
                para.text = f"{chapter_num}.{section_num} {match.group(1)}{match.group(2)}"
                last_level = level
                continue
    now_time_str=time_util.get_now_time_str()
    # D:\bishe\doc_chapter_set
    out_file_path=f"D:\\bishe\\doc_chapter_set\\{now_time_str}.docx"
    print("out_file_path")
    print(out_file_path)
    

    # 保存修改后的文档
    # doc.save(out_file_path)
    # os.system(f'start {out_file_path}')

doc_file=rf"D:\毕设\毕业设计-starp-考试 (修复的).docx"
# doc_file=rf"D:\毕设\毕业设计-starp-考试 (修复的).doc"
apply_heading_style(doc_file)