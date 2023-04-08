import docx
import re

def read_docx(filename):
    # 创建Document对象
    doc = docx.Document(filename)

    # 初始化标题计数器和上一个标题等级
    chapter_num = 0
    last_level = 0
    section_num = 0
    last_section_num = 0

    # 遍历Document对象中的所有Paragraph对象
    for para in doc.paragraphs:
        # 获取当前段落的样式
        style = para.style.name

        # 如果样式是标题，解析标题内容并打印
        if 'Heading' in style:
            level = int(style[-1])
            text = para.text

            # 如果标题以数字开头，表示章节标题
            match = re.match(r'^(\d+)(.*?)$', text.strip())
            if match:
                chapter_num = int(match.group(1))
                section_num = 0
                last_section_num = 0
                print(f"{chapter_num}. {match.group(2)}")
                last_level = level
                continue

            # 如果标题以字母开头，表示小节标题
            match = re.match(r'^([A-Za-z]+)(.*?)$', text.strip())
            if match:
                section_num += 1
                if level > last_level:
                    section_num = last_section_num + 1
                last_section_num = section_num

                # 按照格式打印标题
                print(f"{chapter_num}.{section_num} {match.group(1)}{match.group(2)}")
                last_level = level
                continue


import docx
import re

import time_util
import os

def apply_heading_style(filename):
    # 创建Document对象
    doc = docx.Document(filename)

    # 初始化标题计数器和上一个标题等级
    chapter_num = 0
    last_level = 0
    section_num = 0
    last_section_num = 0

    # 遍历Document对象中的所有Paragraph对象
    for para in doc.paragraphs:
        # 获取当前段落的样式
        style = para.style.name
        # if 'Heading 1' == style: 
        #     # 就 
        #     # if 'Heading 1' == style: 
        
        #     para.style = doc.styles['Heading 1']
        # elif 'Heading 2' == style: 
        #     # 把他设置为 1.1这样子的
        # elif 'Heading 3' == style: 
        #     # 把他设置为 1.1.2这样子的
        # 如果样式是标题，解析标题内容并修改样式
        if 'Heading' in style:
            level = int(style[-1])
            text = para.text
            # 如果他是类似 第1章 绪论，我认为他是一级别的标题
            # 如果标题以数字开头，表示章节标题
            match = re.match(r'^(\d+)(.*?)$', text.strip())
            if match:
                chapter_num = int(match.group(1))
                section_num = 0
                last_section_num = 0
                para.style = doc.styles[f"Heading {level}"]
                para.text = f"{chapter_num}. {match.group(2)}"
                last_level = level
                continue

            # 如果标题以字母开头，表示小节标题
            match = re.match(r'^([A-Za-z]+)(.*?)$', text.strip())
            if match:
                section_num += 1
                if level > last_level:
                    section_num = last_section_num + 1
                last_section_num = section_num

                # 修改样式并设置内容
                para.style = doc.styles[f"Heading {level}"]
                para_text=f"{chapter_num}.{section_num} {match.group(1)}{match.group(2)}"
                print("para_text")
                print(para_text)
                para.text = f"{chapter_num}.{section_num} {match.group(1)}{match.group(2)}"
                last_level = level
                continue
    now_time_str=time_util.get_now_time_str()
    # D:\bishe\doc_chapter_set
    out_file_path=f"D:\\bishe\\doc_chapter_set\\{now_time_str}.docx"
    print("out_file_path")
    print(out_file_path)
    

    # 保存修改后的文档
    # doc.save(out_file_path)
    # os.system(f'start {out_file_path}')


import word_util
from docx.shared import Pt
from docx.enum.style import WD_STYLE

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def check_pic_add_number(para,chapter_num,pic_count):
    # 判断段落中是否有图片
    for run in para.runs:
        if run._element.tag.endswith('}r'):
            # print("r")
            for elem in run._element:
                # print(elem.tag)
                # if elem.tag.endswith('}pict'):
                if  'pict' in elem.tag or 'drawing' in elem.tag:
                    # print("pic")
                    pic_count += 1
                    # 在段落下面插入一个新段落
                    # new_para = para.insert_paragraph_after('')
                    # 图片 居中
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
                    new_para =word_util.insert_paragraph_after(para,'')
                    # new_para =insert_paragraph_after(para,'')
                    # 在新段落中插入自动编号的标签
                    # new_para.add_run 居中
                    pic_run=f"图{chapter_num}.{pic_count}"
                    new_para.add_run(pic_run)
                    # pic_caption_run=new_para.add_run('图{}.{}'.format(pic_count // 10 + 1, pic_count % 10 + 1))
                    # pic_caption_run.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
                    # pic_caption_run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
                    new_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
                    # ————————————————
                    # 版权声明：本文为CSDN博主「Jonathan Star」的原创文章，遵循CC 4.0 BY-SA版权协议，转载请附上原文出处链接及本声明。
                    # 原文链接：https://blog.csdn.net/jonathan_joestar/article/details/129954680
                    # return True
                    # return pic_count
    return pic_count

def apply_heading_style_by_Heading(filename):
    """如果已经是以前面的号码开头了，就没必要在前面加了
    设置 
    一级标题为 宋体 二号
    二级 宋体  小二，
    三级 宋体三号，
    """
    # 创建Document对象
    doc = docx.Document(filename)
    pic_count=0
    

    # 初始化章节计数器和小节计数器
    chapter_num = 0
    section_num = 0

    # 遍历Document对象中的所有Paragraph对象
    for para in doc.paragraphs:
        # 获取当前段落的样式
        style = para.style.name
        text = para.text
        # # 判断段落中是否有图片
        # for run in para.runs:
        #     if run._element.tag.endswith('}r'):
        #         for elem in run._element:
        #             if elem.tag.endswith('}pict'):
        pic_count=check_pic_add_number(para,chapter_num,pic_count)
        # 如果样式是标题，修改样式并设置内容
        
        if 'Heading' in style:
            level = int(style[-1])
            if level == 1:
                chapter_num += 1
                section_num = 0
                # para.style = doc.styles['Heading 1']
                # para.text = f"{chapter_num}. {para.text}"
                # if  text.startswith(f"{chapter_num}.{section_num}"):
                #     continue
                section_subnum=0
                pic_count=0
                para.style.font.name = '宋体'
                para.style.font.size = Pt(22)
            elif level == 2:
                
                section_num += 1
                section_subnum=0
                if  text.startswith(f"{chapter_num}.{section_num}"):
                    continue
                # para.style = doc.styles['Heading 2']
                # if para.text.startswith("第") and para.text.endswith("章"):
                #     para.text = f"{chapter_num}. {para.text}"
                # if para.text.startswith("第") and para.text.endswith("章"):
                #     para.text = f"{chapter_num}. {para.text}"
                para.text = f"{chapter_num}.{section_num} {para.text}"
                para.style.font.name = '宋体'
                para.style.font.size = Pt(16)
                # 标题2 char1 是啥 不要 
            elif level == 3:
                section_subnum+=1
                if  text.startswith(f"{chapter_num}.{section_num}.{section_subnum}"):
                    continue
                # para.style = doc.styles['Heading 3']
                # section_subnum = len([p for p in doc.paragraphs if p.style.name == 'Heading 3'])
                # section_subnum 不是记录前面的所有的，而是从1开始的，如果section_num 变了，他也要从1开始
                para.text = f"{chapter_num}.{section_num}.{section_subnum} {para.text}"
                para.style.font.name = '宋体'
                para.style.font.size = Pt(15)
    now_time_str=time_util.get_now_time_str()
    # D:\bishe\doc_chapter_set
    out_file_path=f"D:\\bishe\\doc_chapter_set\\{now_time_str}.docx"
    print("out_file_path")
    print(out_file_path)


    # 保存修改后的文档
    doc.save(out_file_path)
    os.system(f'start {out_file_path}')

# doc_file=rf"D:\毕设\毕业设计-starp-考试 (修复的).docx"
doc_file=rf"D:\毕设\毕业设计-缪奇鹏-考试 (修复的).docx"
# start "D:\毕设\毕业设计-缪奇鹏-考试 (修复的).docx"

# doc_file=rf"D:\毕设\毕业设计-starp-考试 (修复的).doc"
apply_heading_style_by_Heading(doc_file)

# now_time_str=time_util.get_now_time_str()
# # D:\bishe\doc_chapter_set
# out_file_path=f"D:\\bishe\\doc_chapter_set\\{now_time_str}.docx"
# print("out_file_path")
# print(out_file_path)


# # 保存修改后的文档
# doc.save(out_file_path)
# os.system(f'start {out_file_path}')