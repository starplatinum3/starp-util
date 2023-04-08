import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
input_docx_path="D:\毕业设计-starp-考试.docx"
# "D:\毕业设计-starp-考试.docx"
# input_docx_path='D:\\毕设\\毕业设计-starp-考试.docx'
# 读取Word文档
doc = docx.Document(input_docx_path)

# doc = docx.Document('example.docx')
# "D:\毕业设计-starp-考试.docx"
# "D:\毕设\毕业设计-starp-考试.doc"

# 初始化章节编号和图表编号
chapter_number = 1
figure_number = 1

print("docx.image")
print(docx.image)

import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.image.image import Image
from docx.parts.image import ImagePart
from docx.oxml.shape import CT_Picture

def get_picture(document: Document, paragraph: Paragraph):
    """
    document 为文档对象
    paragraph 为内嵌图片的段落对象
    """
    img = paragraph._element.xpath('.//pic:pic')
    if not img:
        return
    img: CT_Picture = img[0]
    embed = img.xpath('.//a:blip/@r:embed')[0]
    related_part: ImagePart = document.part.related_parts[embed]
    image: Image = related_part.image
    return image

# docx.image. 
# for i in 
# 遍历文档中的每一个段落和图片
# AttributeError: 'Document' object has no attribute 'blocks'
# paragraphs
# doc.blocks:
print("len(doc.paragraphs)")
print(len(doc.paragraphs))
# doc.paragraphs 给图片 加上 标注
for block in doc.paragraphs:
    # print("block._pic")
    # print(block._pic)
    # print("isinstance(block, docx.shape.InlineShape)")
    # print(isinstance(block, docx.shape.InlineShape))
    is_pic_InlineShape=isinstance(block, docx.shape.InlineShape)
    if is_pic_InlineShape:
        print("block._pic")
        print(block._pic)
    # 如果是标题，更新章节编号
    if block.style.name.startswith('Heading'):
        # print("block.style.name")
        # print("block Heading")
        # print(block)
        # print(block.text)
        # block.text 
        chapter_number += 1
        figure_number=1

    # 如果是图片，添加标题和图表编号
    elif isinstance(block, docx.image.Image):
        title = f'图 {chapter_number}.{figure_number}'
        print("block 图")
        print(block)
        # 在图片下方添加图表标题
        paragraph = block.add_paragraph(title, style='Caption')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 更新图表编号
        figure_number += 1
    elif isinstance(block, docx.shape.InlineShape) and block._pic:
        # 创建图表标题
        title = f'图 {chapter_number}.{figure_number}'
        print("block 图")
        print(block)
        # 在图片下方添加图表标题
        paragraph = block.add_paragraph(title, style='Caption')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 更新图表编号
        figure_number += 1

# 保
# D:\毕设\out>
import time_util
now_time_str=time_util.get_now_time_str()

save_path=rf'D:\毕设\out\example_with_captions_{now_time_str}.docx'
print("save_path")
print(save_path)

# 保存修改后的Word文档
# doc.save(save_path)

