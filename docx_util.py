
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
#  python docx 导出 word 表格
def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
 
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

# ModuleNotFoundError: No module named 'docx'

#  pip install python-docx


from docx import Document

doc = Document()
row_cnt=5
col_cnt=4
table = doc.add_table(5, 3, style="Table Grid")
# 字段的个数 是行数 
# 四列  不同的额 
# 遍历 df 
# 获取第1行第3列的单元格（下标从0开始）
cell1 = table.cell(0, 2)
cell1.text = "冰冷的希望"

# df

import pandas as pd
import numpy as np

# df = pd.DataFrame(np.random.randn(4,3),columns = ['col1','col2','col3'])
# for row_index,row in df.iterrows():
#    print(row_index,row)


cell2 = table.cell(1, 2)
paragraph= cell2.paragraphs[0]
run = paragraph.add_run("冰冰很帅")

doc.save('./test_table.docx')
