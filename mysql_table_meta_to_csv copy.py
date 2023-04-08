# select TABLE_NAME,TABLE_TYPE,ENGINE,TABLE_ROWS,TABLE_COMMENT,CREATE_TIME,UPDATE_TIME, CHECK_TIME from information_schema.TABLES where TABLE_SCHEMA='t_shop' ;

"""
select TABLE_NAME,TABLE_TYPE,ENGINE,TABLE_ROWS,TABLE_COMMENT,CREATE_TIME,UPDATE_TIME, CHECK_TIME 
from information_schema.TABLES where TABLE_SCHEMA='t_shop' ;
"""


# table_name='t_exam_paper_answer'
# table_name='t_exam_paper'
# table_name='product'


import pymysql

# D:\proj\private-conf\db_public.json

import json_util
conf=json_util.file_path_to_dict(fr"D:\proj\private-conf\db_public.json")
db_name='public'
SELECT_table_names_sql=f"""
SELECT 

  * ,TABLE_COMMENT

FROM

  information_schema.tables 
  where TABLE_SCHEMA='{db_name}' ;
  """
host=conf['host']
username=conf['username']
password=conf['password']


def select(connection,sql):
    try:
        with connection.cursor() as cursor:
            # sql = '''SHOW TABLES''' # 同理 改成 '''SHOW DATABASES''' 即可获取所有数据库名称
            # sql =f"""
            # select * from information_schema.TABLES where TABLE_SCHEMA='{db_name}' ;

            # """
            # sql = f"""
            # select COLUMN_NAME from information_schema.COLUMNS where table_name = '{table_name}'
            # """
            # t_exam_paper_answer
            # table_name='t_exam_paper_answer'
            # sql = f"""
            # select * from information_schema.COLUMNS WHERE TABLE_SCHEMA = '{db_name}' 
            # """
            cursor.execute(sql)
            table = cursor.fetchall() # 获取所有（输出是列表，列表中的每个元素是字典，字典的KEY是Tables_in_拼接数据库名称）
            for k in range(len(table)): # 提取表名
                # tables.append(table[k]['Tables_in_数据库名称'])
                tables.append(table[k])
    except Exception as e:
        print(e)
    finally:
        connection.commit()
        connection.close()

    return tables

def get_all_tables(db_name,host='localhost',user='root',password="123456"):

    """

    获取数据库中的所有表

    """
    # host='localhost'
    # user='root'
    # password="123456"
    # db="public"
    # db="fastlink"
    # user='用户名',
    connection = pymysql.connect(host=host,
    user=user,

    password=password,

    db=db_name,
    # db=‘数据库名称’,
    charset='utf8',
    cursorclass=pymysql.cursors.DictCursor)

    # connection = pymysql.connect(host=host,

    # user='用户名',

    # password='密码',

    # db='数据库名称',
    # # db=‘数据库名称’,

    # charset='utf8',

    # cursorclass=pymysql.cursors.DictCursor)

    tables = []
    try:
        with connection.cursor() as cursor:
            # sql = '''SHOW TABLES''' # 同理 改成 '''SHOW DATABASES''' 即可获取所有数据库名称
            # sql =f"""
            # select * from information_schema.TABLES where TABLE_SCHEMA='{db_name}' ;

            # """
            # sql = f"""
            # select COLUMN_NAME from information_schema.COLUMNS where table_name = '{table_name}'
            # """
            # t_exam_paper_answer
            # table_name='t_exam_paper_answer'
            sql = f"""
            select * from information_schema.COLUMNS WHERE TABLE_SCHEMA = '{db_name}' 
            """
            cursor.execute(sql)
            table = cursor.fetchall() # 获取所有（输出是列表，列表中的每个元素是字典，字典的KEY是Tables_in_拼接数据库名称）
            for k in range(len(table)): # 提取表名
                # tables.append(table[k]['Tables_in_数据库名称'])
                tables.append(table[k])
    except Exception as e:
        print(e)
    finally:
        connection.commit()
        connection.close()

    return tables


def get_one_table(db_name,table_name,host='localhost',user='root',password="123456"):

    """

    获取数据库中的所有表

    """
    # host='localhost'
    # user='root'
    # password="123456"
    # db="public"
    # db="fastlink"
    # user='用户名',
    connection = pymysql.connect(host=host,
    user=user,

    password=password,

    db=db_name,
    # db=‘数据库名称’,
    charset='utf8',
    cursorclass=pymysql.cursors.DictCursor)

    # connection = pymysql.connect(host=host,

    # user='用户名',

    # password='密码',

    # db='数据库名称',
    # # db=‘数据库名称’,

    # charset='utf8',

    # cursorclass=pymysql.cursors.DictCursor)

    tables = []
    try:
        with connection.cursor() as cursor:
            # sql = '''SHOW TABLES''' # 同理 改成 '''SHOW DATABASES''' 即可获取所有数据库名称
            # sql =f"""
            # select * from information_schema.TABLES where TABLE_SCHEMA='{db_name}' ;

            # """
            # sql = f"""
            # select COLUMN_NAME from information_schema.COLUMNS where table_name = '{table_name}'
            # """
            # t_exam_paper_answer
            # table_name='t_exam_paper_answer'
            sql = f"""
            select * from information_schema.COLUMNS WHERE TABLE_SCHEMA = '{db_name}'  and table_name = '{table_name}'
            """
            cursor.execute(sql)
            table = cursor.fetchall() # 获取所有（输出是列表，列表中的每个元素是字典，字典的KEY是Tables_in_拼接数据库名称）
            for k in range(len(table)): # 提取表名
                # tables.append(table[k]['Tables_in_数据库名称'])
                tables.append(table[k])
    except Exception as e:
        print(e)
    finally:
        connection.commit()
        connection.close()

    return tables

# db_name="fastlink"
# all_tables=get_all_tables(db_name=db_name)

# all_tables=get_all_tables(db_name='public')

# print(all_tables)
tables={
    "physicalTest":"physicalTest",
"physical_test":"physical_test",
"physical_test":"physical_test",
"tChapter":"tChapter",
"chapter":"chapter",
"t_chapter":"t_chapter",
"tExamPaper":"tExamPaper",
"exam_paper":"exam_paper",
"t_exam_paper":"t_exam_paper",
"tExamPaperAnswer":"tExamPaperAnswer",
"exam_paper_answer":"exam_paper_answer",
"t_exam_paper_answer":"t_exam_paper_answer",
"tExamPaperQuestionCustomerAnswer":"tExamPaperQuestionCustomerAnswer",
"exam_paper_question_customer_answer":"exam_paper_question_customer_answer",
"t_exam_paper_question_customer_answer":"t_exam_paper_question_customer_answer",
"tMessage":"tMessage",
"message":"message",
"t_message":"t_message",
"tMessageUser":"tMessageUser",
"message_user":"message_user",
"t_message_user":"t_message_user",
"tQuestion":"tQuestion",
"question":"question",
"t_question":"t_question",
"tQuestion2":"tQuestion2",
"question_2":"question_2",
"t_question_2":"t_question_2",
"tSubject":"tSubject",
"subject":"subject",
"t_subject":"t_subject",
"tTaskExam":"tTaskExam",
"task_exam":"task_exam",
"t_task_exam":"t_task_exam",
"tTaskExamCustomerAnswer":"tTaskExamCustomerAnswer",
"task_exam_customer_answer":"task_exam_customer_answer",
"t_task_exam_customer_answer":"t_task_exam_customer_answer",
"tTextContent":"tTextContent",
"text_content":"text_content",
"t_text_content":"t_text_content",
"tUser":"tUser",
"user":"user",
"t_user":"t_user",
"tUserEventLog":"tUserEventLog",
"user_event_log":"user_event_log",
"t_user_event_log":"t_user_event_log",
"tUserToken":"tUserToken",
"user_token":"user_token",
"t_user_token":"t_user_token",
"tenant":"tenant",
"tenant":"tenant",
"tenant":"tenant",
"tenantExamPaper":"tenantExamPaper",
"tenant_exam_paper":"tenant_exam_paper",
"tenant_exam_paper":"tenant_exam_paper",

}



"""
[{'Tables_in_public': 'component'}, {'Tables_in_public': 'draw'}, {'Tables_in_public': 'physical_test'}, {'Tables_in_public': '
question_draw'}, {'Tables_in_public': 't_chapter'}, {'Tables_in_public': 't_exam_paper'}, {'Tables_in_public': 't_exam_paper_an
swer'}, {'Tables_in_public': 't_exam_paper_question_customer_answer'}, {'Tables_in_public': 't_message'}, {'Tables_in_public':
't_message_user'}, {'Tables_in_public': 't_question'}, {'Tables_in_public': 't_question_2'}, {'Tables_in_public': 't_subject'},
 {'Tables_in_public': 't_task_exam'}, {'Tables_in_public': 't_task_exam_customer_answer'}, {'Tables_in_public': 't_text_content
'}, {'Tables_in_public': 't_user'}, {'Tables_in_public': 't_user_event_log'}, {'Tables_in_public': 't_user_token'}, {'Tables_in
_public': 'tenant'}, {'Tables_in_public': 'tenant_exam_paper'}]
"""

"""
D:\proj\python\my_util_py_pub>python "d:\proj\python\my_util_py_pub\mysql_to_vue_select.py"
[{'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 'component', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB',
'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 3, 'AVG_ROW_LENGTH': 5461, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'I
NDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 3, 'CREATE_TIME': datetime.datetime(2022, 12, 31, 20, 32, 42), 'UPDATE_TIME'
: None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''},
 {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 'draw', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERS
ION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 5, 'AVG_ROW_LENGTH': 3276, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_
LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 7, 'CREATE_TIME': datetime.datetime(2022, 12, 31, 20, 31, 12), 'UPDATE_TIME': Non
e, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TA
BLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 'physical_test', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', '
VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 0, 'AVG_ROW_LENGTH': 0, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX
_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.datetime(2022, 12, 21, 20, 48, 45), 'UPDATE_TIME':
 None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_0900_ai_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''
}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 'question_draw', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'Inn
oDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 0, 'AVG_ROW_LENGTH': 0, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0,
 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 3, 'CREATE_TIME': datetime.datetime(2022, 12, 31, 20, 49, 18), 'UPDATE_TI
ME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_0900_ai_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT'
: ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_chapter', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'Inn
oDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 11, 'AVG_ROW_LENGTH': 1489, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH'
: 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 11, 'CREATE_TIME': datetime.datetime(2022, 12, 21, 20, 48, 47), 'UPDA
TE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COM
MENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_exam_paper', 'TABLE_TYPE': 'BASE TABLE', 'ENGIN
E': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 85, 'AVG_ROW_LENGTH': 192, 'DATA_LENGTH': 16384, 'MAX_DATA_
LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 97, 'CREATE_TIME': datetime.datetime(2022, 12, 21, 20, 48, 47)
, 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TA
BLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_exam_paper_answer', 'TABLE_TYPE': 'BASE
TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 199, 'AVG_ROW_LENGTH': 246, 'DATA_LENGTH': 49
152, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 211, 'CREATE_TIME': datetime.datetime(2022, 12,
 21, 20, 48, 47), 'UPDATE_TIME': datetime.datetime(2023, 1, 19, 13, 51, 57), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_ge
neral_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TA
BLE_NAME': 't_exam_paper_question_customer_answer', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT'
: 'Dynamic', 'TABLE_ROWS': 402, 'AVG_ROW_LENGTH': 163, 'DATA_LENGTH': 65536, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FRE
E': 0, 'AUTO_INCREMENT': 451, 'CREATE_TIME': datetime.datetime(2023, 1, 3, 20, 20, 49), 'UPDATE_TIME': datetime.datetime(2023,
1, 19, 13, 51, 57), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE
_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_message', 'TABLE_TYPE': 'BASE TABLE', 'ENGI
NE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 5, 'AVG_ROW_LENGTH': 3276, 'DATA_LENGTH': 16384, 'MAX_DATA
_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 6, 'CREATE_TIME': datetime.datetime(2022, 12, 21, 20, 48, 47)
, 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TA
BLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_message_user', 'TABLE_TYPE': 'BASE TABLE
', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 14, 'AVG_ROW_LENGTH': 1170, 'DATA_LENGTH': 16384,
'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 15, 'CREATE_TIME': datetime.datetime(2022, 12, 21, 2
0, 48, 47), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS
': '', 'TABLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_question', 'TABLE_TYPE': 'BASE
 TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 77, 'AVG_ROW_LENGTH': 212, 'DATA_LENGTH': 16
384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 131, 'CREATE_TIME': datetime.datetime(2023, 1,
16, 21, 28, 26), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OP
TIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_question_2', 'TABLE_TYPE'
: 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 33, 'AVG_ROW_LENGTH': 496, 'DATA_LENG
TH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 50, 'CREATE_TIME': datetime.datetime(202
2, 12, 21, 20, 48, 47), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CR
EATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_subject', 'TABLE_T
YPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 13, 'AVG_ROW_LENGTH': 1260, 'DATA
_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 16, 'CREATE_TIME': datetime.datetim
e(2022, 12, 21, 20, 48, 47), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None
, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_task_exam', '
TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 0, 'AVG_ROW_LENGTH': 0, 'D
ATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 2, 'CREATE_TIME': datetime.datet
ime(2022, 12, 21, 20, 48, 47), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': No
ne, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_task_exam_c
ustomer_answer', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 0, 'AVG_
ROW_LENGTH': 0, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 1, 'CREATE_TIM
E': datetime.datetime(2022, 12, 21, 20, 48, 47), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_c
i', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAM
E': 't_text_content', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 379
, 'AVG_ROW_LENGTH': 345, 'DATA_LENGTH': 131072, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 16384, 'DATA_FREE': 0, 'AUTO_INCREMENT':
461, 'CREATE_TIME': datetime.datetime(2023, 1, 17, 22, 28, 51), 'UPDATE_TIME': datetime.datetime(2023, 1, 19, 11, 3, 40), 'CHEC
K_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_C
ATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_user', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10
, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 19, 'AVG_ROW_LENGTH': 862, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH':
 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 21, 'CREATE_TIME': datetime.datetime(2022, 12, 21, 23, 6, 42), 'UPDATE_TIME': None, 'CHEC
K_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_C
ATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_user_event_log', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VE
RSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 527, 'AVG_ROW_LENGTH': 155, 'DATA_LENGTH': 81920, 'MAX_DATA_LENGTH': 0, 'IND
EX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 792, 'CREATE_TIME': datetime.datetime(2022, 12, 21, 20, 48, 47), 'UPDATE_TIME'
: datetime.datetime(2023, 1, 19, 13, 51, 57), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'C
REATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 't_user_token', 'TAB
LE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 19, 'AVG_ROW_LENGTH': 862, 'D
ATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 23, 'CREATE_TIME': datetime.date
time(2022, 12, 21, 20, 48, 47), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': N
one, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 'tenant', 'TA
BLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 0, 'AVG_ROW_LENGTH': 0, 'DAT
A_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.date
time(2022, 12, 21, 20, 48, 47), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': N
one, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 'tenant_exam_
paper', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 0, 'AVG_ROW_LENGT
H': 0, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': da
tetime.datetime(2022, 12, 21, 20, 48, 47), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'C
HECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': ''}]
"""

# print("all_tables[0]")
# print(all_tables[0])
# print("table_name")
csharp_type_map={
    'int':'int?',
    'varchar':'string',
    'datetime':'DateTime?',
    'text':'string',
    'decimal':'decimal?',
    'bigint':'long?',
    'tinyint':'bool?'

}
def to_csharp_type(DATA_TYPE):
    if DATA_TYPE in csharp_type_map:
        return csharp_type_map[DATA_TYPE]
    return 'string'
  


import string_util
# print("table_name")
# print(table_name)

# try:
#     print()
# except NameError as e:
#     print(e)

def genElInput():

    for col_info in all_tables:
        COLUMN_NAME=col_info['COLUMN_NAME']
        # to_lower_camel_case()
        col_info['lower_camel_case_COLUMN_NAME']=string_util.to_lower_camel_case(COLUMN_NAME)
        col_info['upper_camel_case_COLUMN_NAME']=string_util.to_upper_camel_case(COLUMN_NAME)
        lower_camel_case_COLUMN_NAME=col_info['lower_camel_case_COLUMN_NAME']
        # input_div=f'<el-input v-model="query.{COLUMN_NAME}"></el-input>'
        # queryParam
        # query
        
        COLUMN_COMMENT=col_info['COLUMN_COMMENT']
        COLUMN_COMMENT_show=COLUMN_COMMENT if COLUMN_COMMENT else COLUMN_NAME
        input_div=f"""
        
         <el-form-item label="{COLUMN_COMMENT_show}">
          <el-input v-model="queryParam.{lower_camel_case_COLUMN_NAME}" clearable  placeholder="{COLUMN_COMMENT_show}"></el-input>
        </el-form-item>
        """
        # input_div=f'{lower_camel_case_COLUMN_NAME}  <el-input v-model="queryParam.{lower_camel_case_COLUMN_NAME}"></el-input>'
        print(input_div)
        # tables.t
"""
{'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'public', 'TABLE_NAME': 'component', 'COLUMN_NAME': 'id', 'ORDINAL_POSITION': 1, 'COLU
MN_DEFAULT': None, 'IS_NULLABLE': 'NO', 'DATA_TYPE': 'int', 'CHARACTER_MAXIMUM_LENGTH': None, 'CHARACTER_OCTET_LENGTH': None, '
NUMERIC_PRECISION': 10, 'NUMERIC_SCALE': 0, 'DATETIME_PRECISION': None, 'CHARACTER_SET_NAME': None, 'COLLATION_NAME': None, 'CO
LUMN_TYPE': 'int', 'COLUMN_KEY': 'PRI', 'EXTRA': 'auto_increment', 'PRIVILEGES': 'select,insert,update,references', 'COLUMN_COMMENT': '编号', 'GENERATION_EXPRESSION': '', 'SRS_ID': None}
"""

def genCSharpClass():
    print(all_tables)
    # table_name
    table_name_class_name=string_util.to_upper_camel_case(table_name)
    # print(f"public class Blog {")
    row_val_list=[]
    for col_info in all_tables:
        COLUMN_NAME=col_info['COLUMN_NAME']
        # to_lower_camel_case()
        col_info['lower_camel_case_COLUMN_NAME']=string_util.to_lower_camel_case(COLUMN_NAME)
        col_info['upper_camel_case_COLUMN_NAME']=string_util.to_upper_camel_case(COLUMN_NAME)
        lower_camel_case_COLUMN_NAME=col_info['lower_camel_case_COLUMN_NAME']
        # input_div=f'<el-input v-model="query.{COLUMN_NAME}"></el-input>'
        # queryParam
        # query
        
        COLUMN_COMMENT=col_info['COLUMN_COMMENT']
        COLUMN_COMMENT_show=COLUMN_COMMENT if COLUMN_COMMENT else COLUMN_NAME
        # public class Blog {
        #     [Column(IsIdentity = true, IsPrimary = true)]
        #     public int BlogId { get; set; }
        #     public string Url { get; set; }
        #     public int Rating { get; set; }
        # }
        DATA_TYPE=col_info['DATA_TYPE']
        csharp_type=to_csharp_type(DATA_TYPE)
        # input_div= "public int {COLUMN_COMMENT} { get; set; }".replace("{COLUMN_COMMENT}",COLUMN_COMMENT)
        input_div= "public {csharp_type} {COLUMN_NAME} { get; set; }".replace("{COLUMN_NAME}",COLUMN_NAME)\
        .replace("{DATA_TYPE}",DATA_TYPE)\
        .replace("{csharp_type}",csharp_type)\
        # input_div=f"""
        #  public int COLUMN_COMMENT { get; set; }
        row_val_list.append(input_div)
        # """
        # input_div=f'{lower_camel_case_COLUMN_NAME}  <el-input v-model="queryParam.{lower_camel_case_COLUMN_NAME}"></el-input>'
        # print(input_div)
        # tables.t
    # row_val_list.join()
    field_rows="\n".join(row_val_list)
    class_file_tpl="""
        using FreeSql.DataAnnotations;
        using System;

        public class {table_name_class_name} {
            [Column(IsIdentity = true, IsPrimary = true)]
            {field_rows}
        }"""
    return class_file_tpl.replace("{table_name_class_name}",table_name_class_name).\
        replace("{field_rows}",field_rows)

# def toFieldMysqlMapEntity(col_info):

def toFieldMysqlMapEntity(col_info):
    DATA_TYPE=col_info['DATA_TYPE']
    COLUMN_NAME=col_info['COLUMN_NAME']
    # line.{COLUMN_NAME} = {COLUMN_NAME};
    if DATA_TYPE=='int':
        return f"""
        int {COLUMN_NAME} = 0;
        int.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});
        """
    if DATA_TYPE=='datetime':
        return f"""
        DateTime {COLUMN_NAME} = DateTime.Now;
        DateTime.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='decimal':
        return f"""
        decimal {COLUMN_NAME} = 0;
        decimal.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='varchar':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='text':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='tinyint':
        return f"""
        int {COLUMN_NAME} = 0;
        int.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='bigint':
        return f"""
        long {COLUMN_NAME} = 0;
        long.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='float':
        return f"""
        float {COLUMN_NAME} = 0;
        float.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='double':
        return f"""
        double {COLUMN_NAME} = 0;
        double.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='bit':
        return f"""
        bool {COLUMN_NAME} = false;
        bool.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='char':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='date':
        return f"""
        DateTime {COLUMN_NAME} = DateTime.Now;
        DateTime.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='time':
        return f"""
        TimeSpan {COLUMN_NAME} = new TimeSpan();
        TimeSpan.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='timestamp':
        return f"""
        DateTime {COLUMN_NAME} = DateTime.Now;
        DateTime.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='year':
        return f"""
        int {COLUMN_NAME} = 0;
        int.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='mediumint':
        return f"""
        int {COLUMN_NAME} = 0;
        int.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='smallint':
        return f"""
        int {COLUMN_NAME} = 0;
        int.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});"""
    if DATA_TYPE=='tinytext':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='mediumtext':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='longtext':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='enum':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='set':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='blob':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='mediumblob':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='longblob':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='varbinary':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='binary':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='geometry':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='point':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    if DATA_TYPE=='linestring':
        return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""
    return f"""
        string {COLUMN_NAME} = reader["{COLUMN_NAME}"].ToString();"""

#     """
#     int id = 0;
#    int.TryParse(reader["id"].ToString(), out id);"""

def genMysqlMapToEntity():
    print(all_tables)
    # table_name
    table_name_class_name=string_util.to_upper_camel_case(table_name)
    # print(f"public class Blog {")
    row_val_list=[]
    for col_info in all_tables:
        COLUMN_NAME=col_info['COLUMN_NAME']
        # to_lower_camel_case()
        col_info['lower_camel_case_COLUMN_NAME']=string_util.to_lower_camel_case(COLUMN_NAME)
        col_info['upper_camel_case_COLUMN_NAME']=string_util.to_upper_camel_case(COLUMN_NAME)
        lower_camel_case_COLUMN_NAME=col_info['lower_camel_case_COLUMN_NAME']
        # input_div=f'<el-input v-model="query.{COLUMN_NAME}"></el-input>'
        # queryParam
        # query
        
        COLUMN_COMMENT=col_info['COLUMN_COMMENT']
        # DATA_TYPE=col_info['DATA_TYPE']
        COLUMN_COMMENT_show=COLUMN_COMMENT if COLUMN_COMMENT else COLUMN_NAME
        # if DATA_TYPE:
        # public class Blog {
        #     [Column(IsIdentity = true, IsPrimary = true)]
        #     public int BlogId { get; set; }
        #     public string Url { get; set; }
        #     public int Rating { get; set; }
        # }
        # DATA_TYPE
        FieldMysqlMapEntity=toFieldMysqlMapEntity(col_info)
        # DATA_TYPE=col_info['DATA_TYPE']
        # COLUMN_NAME=col_info['COLUMN_NAME']
        
        # if DATA_TYPE=='int':
        #     return f"""
        #     int {COLUMN_NAME} = 0;
        #     int.TryParse(reader["{COLUMN_NAME}"].ToString(), out {COLUMN_NAME});
        #     line.{COLUMN_NAME} = {COLUMN_NAME};
        #     """
        # line.{COLUMN_NAME} = {COLUMN_NAME};
        # input_div= "public int {COLUMN_COMMENT} { get; set; }".replace("{COLUMN_COMMENT}",COLUMN_COMMENT)
        # input_div= "public int {COLUMN_NAME} { get; set; }".replace("{COLUMN_NAME}",COLUMN_NAME)
        # input_div=f"""
        #  public int COLUMN_COMMENT { get; set; }
        row_val_list.append(FieldMysqlMapEntity)
        row_val_list.append(f"line.{COLUMN_NAME} = {COLUMN_NAME};")
        # row_val_list.append(input_div)
        # """
        # input_div=f'{lower_camel_case_COLUMN_NAME}  <el-input v-model="queryParam.{lower_camel_case_COLUMN_NAME}"></el-input>'
        # print(input_div)
        # tables.t
    # row_val_list.join()
    field_rows="\n".join(row_val_list)
    class_file_tpl="""
        using FreeSql.DataAnnotations;
        using System;

        public class {table_name_class_name} {
            [Column(IsIdentity = true, IsPrimary = true)]
            {field_rows}
        }"""
    return class_file_tpl.replace("{table_name_class_name}",table_name_class_name).\
        replace("{field_rows}",field_rows)


def genCreate():
    print(all_tables)
    # table_name
    table_name_class_name=string_util.to_upper_camel_case(table_name)
    # print(f"public class Blog {")
    row_val_list=[]
    COLUMN_NAME_list=[]
    COLUMN_NAME_at_list=[]
    AddWithValue_row_list=[]
    for col_info in all_tables:
        COLUMN_NAME=col_info['COLUMN_NAME']
        COLUMN_NAME_list.append(COLUMN_NAME)
        COLUMN_NAME_at_list.append(f"@{COLUMN_NAME}")
        # to_lower_camel_case()
        col_info['lower_camel_case_COLUMN_NAME']=string_util.to_lower_camel_case(COLUMN_NAME)
        col_info['upper_camel_case_COLUMN_NAME']=string_util.to_upper_camel_case(COLUMN_NAME)
        lower_camel_case_COLUMN_NAME=col_info['lower_camel_case_COLUMN_NAME']
        # input_div=f'<el-input v-model="query.{COLUMN_NAME}"></el-input>'
        # queryParam
        # query
        AddWithValue_row=f'com.Parameters.AddWithValue("@{COLUMN_NAME}", line.{COLUMN_NAME});'
        AddWithValue_row_list.append(AddWithValue_row)
        
        COLUMN_COMMENT=col_info['COLUMN_COMMENT']
        # DATA_TYPE=col_info['DATA_TYPE']
        COLUMN_COMMENT_show=COLUMN_COMMENT if COLUMN_COMMENT else COLUMN_NAME
        # if DATA_TYPE:
        # public class Blog {
        #     [Column(IsIdentity = true, IsPrimary = true)]
        #     public int BlogId { get; set; }
        #     public string Url { get; set; }
        #     public int Rating { get; set; }
        # }
        # DATA_TYPE
        FieldMysqlMapEntity=toFieldMysqlMapEntity(col_info)

        row_val_list.append(FieldMysqlMapEntity)
        row_val_list.append(f"line.{COLUMN_NAME} = {COLUMN_NAME};")

    field_rows="\n".join(row_val_list)
    

    COLUMN_NAMEs=", ".join(COLUMN_NAME_list)
    COLUMN_NAME_at_list_str=", ".join(COLUMN_NAME_at_list)
    AddWithValue_row_list_str="\n".join(AddWithValue_row_list)
    class_file_tpl="""
        //create
        [HttpPost]
        [Route("api/{table_name}")]
        // POST: api/{table_name}
        public void Post([FromBody]{table_name_class_name} line)
        {
            using (MySqlConnection conn = this.GetConnection())
            {
                using (MySqlCommand com = new MySqlCommand())
                {
                    com.Connection = conn;
                    com.CommandText = "insert into {table_name}({COLUMN_NAMEs}) values ({COLUMN_NAME_at_list_str}) ";
                    {AddWithValue_row_list_str}
                    com.ExecuteNonQuery();
                }
            }
        }
        """
    return class_file_tpl.replace("{table_name_class_name}",table_name_class_name).\
        replace("{field_rows}",field_rows)\
        .replace("{COLUMN_NAMEs}",COLUMN_NAMEs)\
        .replace("{COLUMN_NAME_at_list_str}",COLUMN_NAME_at_list_str)\
        .replace("{AddWithValue_row_list_str}",AddWithValue_row_list_str)\
        .replace("{table_name}",table_name)\




def genPut():
    print(all_tables)
    # table_name
    table_name_class_name=string_util.to_upper_camel_case(table_name)
    # print(f"public class Blog {")
    row_val_list=[]
    COLUMN_NAME_list=[]
    COLUMN_NAME_at_list=[]
    AddWithValue_row_list=[]
    COLUMN_NAME_update_set_list=[]
    if_not_null_select_list=[]
    for col_info in all_tables:
        COLUMN_NAME=col_info['COLUMN_NAME']
        COLUMN_NAME_list.append(COLUMN_NAME)
        COLUMN_NAME_at_list.append(f"@{COLUMN_NAME}")
        if COLUMN_NAME!="id":
            COLUMN_NAME_update_set_list.append(f"{COLUMN_NAME} = @{COLUMN_NAME}")
        # birthday=@birthday
        # to_lower_camel_case()
        
        if_tpl="""
        if (postData.{COLUMN_NAME} != null)
        {
            values.Add($" {COLUMN_NAME} = @{COLUMN_NAME} ");
            com.Parameters.AddWithValue("@{COLUMN_NAME}", postData.{COLUMN_NAME});
        }
        """
        if_not_null_select=if_tpl.replace("{COLUMN_NAME}",COLUMN_NAME)
        if_not_null_select_list.append(if_not_null_select)
        col_info['lower_camel_case_COLUMN_NAME']=string_util.to_lower_camel_case(COLUMN_NAME)
        col_info['upper_camel_case_COLUMN_NAME']=string_util.to_upper_camel_case(COLUMN_NAME)
        lower_camel_case_COLUMN_NAME=col_info['lower_camel_case_COLUMN_NAME']
        # input_div=f'<el-input v-model="query.{COLUMN_NAME}"></el-input>'
        # queryParam
        # query
        AddWithValue_row=f'com.Parameters.AddWithValue("@{COLUMN_NAME}", line.{COLUMN_NAME});'
        AddWithValue_row_list.append(AddWithValue_row)
        
        COLUMN_COMMENT=col_info['COLUMN_COMMENT']
        # DATA_TYPE=col_info['DATA_TYPE']
        COLUMN_COMMENT_show=COLUMN_COMMENT if COLUMN_COMMENT else COLUMN_NAME
        # if DATA_TYPE:
        # public class Blog {
        #     [Column(IsIdentity = true, IsPrimary = true)]
        #     public int BlogId { get; set; }
        #     public string Url { get; set; }
        #     public int Rating { get; set; }
        # }
        # DATA_TYPE
        FieldMysqlMapEntity=toFieldMysqlMapEntity(col_info)

        row_val_list.append(FieldMysqlMapEntity)
        row_val_list.append(f"line.{COLUMN_NAME} = {COLUMN_NAME};")

    field_rows="\n".join(row_val_list)
    

    COLUMN_NAMEs=", ".join(COLUMN_NAME_list)
    COLUMN_NAME_at_list_str=", ".join(COLUMN_NAME_at_list)
    AddWithValue_row_list_str="\n".join(AddWithValue_row_list)
    COLUMN_NAME_update_set_list_str=", ".join(COLUMN_NAME_update_set_list)
    if_not_null_select_list_str="\n".join(if_not_null_select_list)
    class_file_tpl="""
        [HttpPost]
        [Route("api/{table_name}/put")]
        // PUT: api/{table_name}/5
        public object Put( [FromBody] {table_name_class_name} line)
        {
            using (MySqlConnection conn = this.GetConnection())
            {
                using (MySqlCommand com = new MySqlCommand())
                {
                    com.Connection = conn;
                    com.CommandText = "update {table_name} set {COLUMN_NAME_update_set_list_str} where id=@id ";
                    {AddWithValue_row_list_str}
                    com.ExecuteNonQuery();
                }
            }
             return ReturnT.Ok("ok");
        }

        [HttpPost]
        [Route("api/{table_name}/delete")]
        // DELETE: api/{table_name}/5
        public object Delete([FromBody] {table_name_class_name} line)
        {
            using (MySqlConnection conn = this.GetConnection())
            {
                using (MySqlCommand com = new MySqlCommand())
                {
                    com.Connection = conn;
                    com.CommandText = "delete from {table_name} where id=@id ";
                    com.Parameters.AddWithValue("@id", line.id);
                    com.ExecuteNonQuery();
                }
            }
             return ReturnT.Ok("ok");
        }


         [HttpPost]
        [Route("api/{table_name}")]
        // GET api/{table_name}
        public object selectPage([FromBody] {table_name_class_name}  postData,int pageIndex,int pageSize)
        {
            List<{table_name_class_name}> list = new List<{table_name_class_name}>();
            using (MySqlConnection conn = this.GetConnection())
            {
                using (MySqlCommand com = new MySqlCommand())
                {
                    com.Connection = conn;
                    com.CommandText = "select * from {table_name} where ";
                    MySqlDataReader reader = com.ExecuteReader();
                    while (reader.Read())
                    {
                        {table_name_class_name} line   = new {table_name_class_name}();
                       {field_rows}
                    }
                }
            }
            return list.ToArray();
        }

        [HttpPost]
           [Route("api/{table_name}/selectPage")]
        // GET api/{table_name}
       public object selectPage([FromBody] {table_name_class_name}  postData,int pageIndex,int pageSize)
        {
             List<{table_name_class_name}> list = new List<{table_name_class_name}>();
            using (MySqlConnection conn = this.GetConnection())
            {
                using (MySqlCommand com = new MySqlCommand())
                {
                    com.Connection = conn;
                    string sql = "select * from {table_name}   ";

                    List<string> values = new List<string>();
                    {if_not_null_select_list_str}
                  
                  
                    string whereCondition = string.Join(" \\n and ", values);
           
                    if (values.Count > 0)
                    {
                        sql +="  where  "+ whereCondition;
                    }

                    com.CommandText = sql;
                   
                    MySqlDataReader reader = com.ExecuteReader();
                    while (reader.Read())
                    {
                         {table_name_class_name} line   = new {table_name_class_name}(); 
                       {field_rows}
                        list.Add(line);
                    }
                }
            }
            return list.ToArray();
        }

        """
    return class_file_tpl.replace("{table_name_class_name}",table_name_class_name).\
        replace("{field_rows}",field_rows)\
        .replace("{COLUMN_NAMEs}",COLUMN_NAMEs)\
        .replace("{COLUMN_NAME_at_list_str}",COLUMN_NAME_at_list_str)\
        .replace("{AddWithValue_row_list_str}",AddWithValue_row_list_str)\
        .replace("{table_name}",table_name)\
        .replace("{COLUMN_NAME_update_set_list_str}",COLUMN_NAME_update_set_list_str)\
        .replace("{if_not_null_select_list_str}",if_not_null_select_list_str)\
        

def genLunWenDesp():
    
    db_name="public"
    table_name="java"
    # sql 获取所有的列名
    # table_name
    all_tables=get_all_tables(db_name=db_name,\
                              host=conf['host'],user=conf['username'],password=conf['password'])
    
    print("all_tables")
    print(all_tables)
    table_name_class_name=string_util.to_upper_camel_case(table_name)
    # print(f"public class Blog {")
    row_val_list=[]
    COLUMN_NAME_list=[]
    COLUMN_NAME_at_list=[]
    AddWithValue_row_list=[]
    COLUMN_NAME_update_set_list=[]
    if_not_null_select_list=[]
    for col_info in all_tables:
        COLUMN_NAME=col_info['COLUMN_NAME']
        COLUMN_NAME_list.append(COLUMN_NAME)
        COLUMN_NAME_at_list.append(f"@{COLUMN_NAME}")
        if COLUMN_NAME!="id":
            COLUMN_NAME_update_set_list.append(f"{COLUMN_NAME} = @{COLUMN_NAME}")
        # birthday=@birthday
        # to_lower_camel_case()
        
        if_tpl="""
        if (postData.{COLUMN_NAME} != null)
        {
            values.Add($" {COLUMN_NAME} = @{COLUMN_NAME} ");
            com.Parameters.AddWithValue("@{COLUMN_NAME}", postData.{COLUMN_NAME});
        }
        """
        if_not_null_select=if_tpl.replace("{COLUMN_NAME}",COLUMN_NAME)
        if_not_null_select_list.append(if_not_null_select)
        col_info['lower_camel_case_COLUMN_NAME']=string_util.to_lower_camel_case(COLUMN_NAME)
        col_info['upper_camel_case_COLUMN_NAME']=string_util.to_upper_camel_case(COLUMN_NAME)
        lower_camel_case_COLUMN_NAME=col_info['lower_camel_case_COLUMN_NAME']
        # input_div=f'<el-input v-model="query.{COLUMN_NAME}"></el-input>'
        # queryParam
        # query
        AddWithValue_row=f'com.Parameters.AddWithValue("@{COLUMN_NAME}", line.{COLUMN_NAME});'
        AddWithValue_row_list.append(AddWithValue_row)
        
        COLUMN_COMMENT=col_info['COLUMN_COMMENT']
        # DATA_TYPE=col_info['DATA_TYPE']
        COLUMN_COMMENT_show=COLUMN_COMMENT if COLUMN_COMMENT else COLUMN_NAME
        # if DATA_TYPE:
        # public class Blog {
        #     [Column(IsIdentity = true, IsPrimary = true)]
        #     public int BlogId { get; set; }
        #     public string Url { get; set; }
        #     public int Rating { get; set; }
        # }
        # DATA_TYPE
        FieldMysqlMapEntity=toFieldMysqlMapEntity(col_info)

        row_val_list.append(FieldMysqlMapEntity)
        row_val_list.append(f"line.{COLUMN_NAME} = {COLUMN_NAME};")

    field_rows="\n".join(row_val_list)
    

    COLUMN_NAMEs=", ".join(COLUMN_NAME_list)
    COLUMN_NAME_at_list_str=", ".join(COLUMN_NAME_at_list)
    AddWithValue_row_list_str="\n".join(AddWithValue_row_list)
    COLUMN_NAME_update_set_list_str=", ".join(COLUMN_NAME_update_set_list)
    if_not_null_select_list_str="\n".join(if_not_null_select_list)
    class_file_tpl="""
        [HttpPost]
        [Route("api/{table_name}/put")]
        // PUT: api/{table_name}/5
        public object Put( [FromBody] {table_name_class_name} line)
        {
            using (MySqlConnection conn = this.GetConnection())
            {
                using (MySqlCommand com = new MySqlCommand())
                {
                    com.Connection = conn;
                    com.CommandText = "update {table_name} set {COLUMN_NAME_update_set_list_str} where id=@id ";
                    {AddWithValue_row_list_str}
                    com.ExecuteNonQuery();
                }
            }
             return ReturnT.Ok("ok");
        }

        [HttpPost]
        [Route("api/{table_name}/delete")]
        // DELETE: api/{table_name}/5
        public object Delete([FromBody] {table_name_class_name} line)
        {
            using (MySqlConnection conn = this.GetConnection())
            {
                using (MySqlCommand com = new MySqlCommand())
                {
                    com.Connection = conn;
                    com.CommandText = "delete from {table_name} where id=@id ";
                    com.Parameters.AddWithValue("@id", line.id);
                    com.ExecuteNonQuery();
                }
            }
             return ReturnT.Ok("ok");
        }


         [HttpPost]
        [Route("api/{table_name}")]
        // GET api/{table_name}
        public object selectPage([FromBody] {table_name_class_name}  postData,int pageIndex,int pageSize)
        {
            List<{table_name_class_name}> list = new List<{table_name_class_name}>();
            using (MySqlConnection conn = this.GetConnection())
            {
                using (MySqlCommand com = new MySqlCommand())
                {
                    com.Connection = conn;
                    com.CommandText = "select * from {table_name} where ";
                    MySqlDataReader reader = com.ExecuteReader();
                    while (reader.Read())
                    {
                        {table_name_class_name} line   = new {table_name_class_name}();
                       {field_rows}
                    }
                }
            }
            return list.ToArray();
        }

        [HttpPost]
           [Route("api/{table_name}/selectPage")]
        // GET api/{table_name}
       public object selectPage([FromBody] {table_name_class_name}  postData,int pageIndex,int pageSize)
        {
             List<{table_name_class_name}> list = new List<{table_name_class_name}>();
            using (MySqlConnection conn = this.GetConnection())
            {
                using (MySqlCommand com = new MySqlCommand())
                {
                    com.Connection = conn;
                    string sql = "select * from {table_name}   ";

                    List<string> values = new List<string>();
                    {if_not_null_select_list_str}
                  
                  
                    string whereCondition = string.Join(" \\n and ", values);
           
                    if (values.Count > 0)
                    {
                        sql +="  where  "+ whereCondition;
                    }

                    com.CommandText = sql;
                   
                    MySqlDataReader reader = com.ExecuteReader();
                    while (reader.Read())
                    {
                         {table_name_class_name} line   = new {table_name_class_name}(); 
                       {field_rows}
                        list.Add(line);
                    }
                }
            }
            return list.ToArray();
        }

        """
    return class_file_tpl.replace("{table_name_class_name}",table_name_class_name).\
        replace("{field_rows}",field_rows)\
        .replace("{COLUMN_NAMEs}",COLUMN_NAMEs)\
        .replace("{COLUMN_NAME_at_list_str}",COLUMN_NAME_at_list_str)\
        .replace("{AddWithValue_row_list_str}",AddWithValue_row_list_str)\
        .replace("{table_name}",table_name)\
        .replace("{COLUMN_NAME_update_set_list_str}",COLUMN_NAME_update_set_list_str)\
        .replace("{if_not_null_select_list_str}",if_not_null_select_list_str)\
        

# COLUMN_NAME

# genElInput()
   
# CSharpClass=genCSharpClass()

# print(CSharpClass)
"""

  public class Product {
            [Column(IsIdentity = true, IsPrimary = true)]
            public int create_time { get; set; }
public int id { get; set; }
public int name { get; set; }
public int product_code { get; set; }
public int specification { get; set; }
        }
        """

import time_util

def gen_test1():
    table_name='product'
    print("table_name")
    print(table_name)
    db_name="fastlink"
    all_tables=get_all_tables(db_name=db_name)
    # D:\codeGen\genCSharp
    genOutDir=rf"D:\codeGen\genCSharp"
    # res=genMysqlMapToEntity()
    # res=genCSharpClass()
    # res=genCreate()
    # # genName="genCSharpClass"
    # genName="genCreate"

    res=genPut()
    genName="genPut"

    print(res)

    
    now_time_str=time_util.get_now_time_str()

    # import
    # def get_now_time_str():
    #     now_time_str = time.strftime("%Y_%m_%d_%H_%M_%S", time.localtime())
    #     return now_time_str


    # get_now_time_str=get_now_time()


    outFileName=f"{genOutDir}/{table_name}_{genName}_{now_time_str}.cs"
    print("outFileName")
    print(outFileName)

    with open(outFileName, "w") as f:
        f.write(res)

    # with open(f"{table_name}_genMysqlMapToEntity.cs", "w") as f:
    #     f.write(res)

import file_util
import pandas_util
import pandas as pd

def df_chinese_name_set(df):
    df['类型']=df['COLUMN_TYPE']
    # COLUMN_TYPE
    df['是否为空']=df['IS_NULLABLE']
    df['键']=df['COLUMN_KEY']
    df['备注']=df['COLUMN_COMMENT']
    df['字段']=df['COLUMN_NAME']
    
    return  df


import time_util

def get_table_comment(tables,need_name):
    # pri 
    print("need_name")
    print(need_name)

    for table in tables:
        if table['TABLE_NAME']==need_name:
            print("table")
            print(table)
            TABLE_COMMENT=table.get('TABLE_COMMENT')
            if TABLE_COMMENT is None:
                return need_name
            return TABLE_COMMENT
    return None
    
from docx import Document

import docx
print("docx.__version__")
print(docx.__version__)

import logging
from docx.enum.style import WD_STYLE

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# document = Document()
# table = document.add_table(rows=3, cols=3)

# 添加一个段落用于存放表格题注
# caption_para = document.add_paragraph()


def mysql_table_meta_to_csv():
    do_write=False
    table_name='component'
    print("table_name")
    print(table_name)
    # db_name="fastlink"
    db_name="public"
    # csv 中文  乱码 
    # COLUMN_NAME
    connection = pymysql.connect(host=host,
    user=username,
    password=password,
    db=db_name,
    # db=‘数据库名称’,
    charset='utf8',
    cursorclass=pymysql.cursors.DictCursor)
    # tables=[]
    table_number=1
    """
    {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'all_visit_count', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 31, 'AVG_ROW_LENGTH': 528, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 56), 'UPDATE_TIME': datetime.datetime(2022, 12, 13, 10, 3, 26), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'correct_count', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 10, 'AVG_ROW_LENGTH': 1638, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 57), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'student', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 48, 'AVG_ROW_LENGTH': 341, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 217, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 57), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'visit_time', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 140, 'AVG_ROW_LENGTH': 117, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 57), 'UPDATE_TIME': datetime.datetime(2022, 12, 13, 10, 3, 26), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'mqtt_control', 'TABLE_NAME': 'backdirlog', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 25, 'AVG_ROW_LENGTH': 655, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 35, 'CREATE_TIME': datetime.datetime(2022, 6, 9, 14, 57, 40), 'UPDATE_TIME': datetime.datetime(2022, 11, 19, 9, 37, 28), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'mqtt_control', 'TABLE_NAME': 'git_down', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 4, 'AVG_ROW_LENGTH': 4096, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 5, 'CREATE_TIME': datetime.datetime(2022, 6, 9, 21, 56, 42), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}]"""
    with connection.cursor() as cursor:
        cursor.execute(SELECT_table_names_sql)
        tables=cursor.fetchall()
        print("tables")
        print(tables)
    all_tables=get_all_tables(db_name=db_name,\
                              host=conf['host'],user=conf['username'],password=conf['password'])
    # print(all_tables)
    # out_file_dir=rf"D:\毕设"
    # D:\毕设\colInfo
    doc = Document()
    now_time_str=time_util.get_now_time_str()
    out_file_dir=rf"D:\毕设\colInfo/{now_time_str}"
    if do_write:
        os.makedirs(out_file_dir)
    # out_file_name=f"{out_file_dir}/{now_time_str}/{db_name}.csv"
    # print("out_file_name")
    # print(out_file_name)
    # file_util.to_csv(all_tables,out_file_name)
    # 管理员表包括主要七个字段，主要定义如表3.1所示。
    df=pd.DataFrame.from_records(all_tables)
    # df.columns
    # columns_len=len(df.columns)
    
    # df 只要某几列 
    # df  不同的 tablename 写入不同的 csv文件
    # need_columns=['COLUMN_NAME','DATA_TYPE','COLUMN_COMMENT','IS_NULLABLE','COLUMN_KEY','TABLE_NAME']
    need_columns=['字段','类型','是否为空','备注','键','COLUMN_NAME','DATA_TYPE','COLUMN_COMMENT','IS_NULLABLE','COLUMN_KEY','TABLE_NAME']
    # df['类型']=df['DATA_TYPE']
    col_cnt=4
    # df groupby tablename 
    # data.groupby("company")
    # df['TABLE_NAME']
    TABLE_NAME_set=set(df['TABLE_NAME'])
    # df=df_chinese_name_set(df)
    df=df_chinese_name_set(df)
    table_info_list=[]
    
    # df = pd.DataFrame(np.random.randn(4,3),columns = ['col1','col2','col3'])
    # for row_index,row in df.iterrows():
    #    print(row_index,row)
    # 根据上述的相关需求，系统需要管理员表、会议室表、设备表、用户表、历史记录表、用户管理会议室多对多关系表这6张表，各表的详细设计和相关关联关系如图3.2所示。
    TABLE_NAME_says="表、".join(TABLE_NAME_set)

    for TABLE_NAME in TABLE_NAME_set:
        # print(TABLE_NAME)
        df_table=df[df['TABLE_NAME']==TABLE_NAME]
        df_table=df_table[need_columns]
        print(df_table) 
        # df 计算 行数
        row_cnt=len(df_table)
        # doc 加入一行 python 
        # AttributeError: 'Table' object has no attribute 'add_caption'
        table_comment=get_table_comment(tables,TABLE_NAME)
        

        table_info=f"{table_comment}表包括主要{row_cnt}个字段，主要定义如表3.{table_number}所示。"
        paragraph1 = doc.add_paragraph(table_info)
        pandas_util.docx_add_df_table(df_table,doc,need_columns=['字段', '类型', '是否为空', '备注'])
        # 表加上 题注 python doc 
        document=doc
        # 添加一个段落用于存放表格题注
        caption_para = document.add_paragraph()
        # ValueError: assigned style is type PARAGRAPH (1), need type CHARACTER (2)
        caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加一个 run 用于存放题注内容
        caption_run = caption_para.add_run()
        caption_run.text = f"表3.{table_number}"
        
        # 将该 run 的样式设置为题注样式（需要自行定义样式或者使用默认样式）
        # caption_run.style = document.styles["Caption"]
        # styles = document.styles
        # document.styles
        print("document.styles[")
        print(document.styles)
        # caption_run.style = document.styles[WD_STYLE.CAPTION]
        # KeyError: "no style with name 'CAPTION (-35)'"
        # for style in document.styles:
        #     print(style.name)
        """
            Normal
        Header
        Header Char
        Footer
        Footer Char
        Heading 1
        Heading 2
        Heading 3
        Heading 4
        Heading 5
        Heading 6
        """
        # print(f"管理员表包括主要{row_cnt}个字段，主要定义如表3.1所示。")
        
        table_info_list.append(table_info)
        print(f"{table_comment}表包括主要{row_cnt}个字段，主要定义如表3.{table_number}所示。")
        table_number+=1
        # out_table_file_path=f"{out_file_dir}/{now_time_str}/{TABLE_NAME}.csv"
        # pandas_util.to_csv(df_table,f"{out_file_dir}/{TABLE_NAME}.csv")
        out_file_name=f"{out_file_dir}/{TABLE_NAME}.csv"
        if do_write:
            pandas_util.to_csv(df_table,out_file_name)
        print("out_file_name")
        print(out_file_name)
    # table_info_list
    # table_info_list. 
    table_info_doc="\n".join(table_info_list)
    print("table_info_doc")
    print(table_info_doc)
    # python 生成word 表格 有格式的线
    print("TABLE_NAME_says")
    print(TABLE_NAME_says)
    if do_write:
        doc.save('./test_tables_db.docx')



def str_is_none(string):
    if string is None:
        return True
    if string=="":
        return True
    return False

def get_TABLE_COMMENT_show(table_obj):
    
    TABLE_COMMENT=table_obj['TABLE_COMMENT']
    if str_is_none(TABLE_COMMENT):
        table_name=table_obj['TABLE_NAME']
        return table_name
    return TABLE_COMMENT

def genTableLunWen(table_obj,cols,db_name="public"):
    pass
    table_name=table_obj['TABLE_NAME']
    TABLE_COMMENT=table_obj['TABLE_COMMENT']
    TABLE_COMMENT_show=get_TABLE_COMMENT_show(table_obj)
    print("table_obj")
    print(table_obj)
    # TABLE_COMMENT
    # db_name="public"
    # table_name="java"
    # sql 获取所有的列名
    # table_name
    # all_tables=get_all_tables(db_name=db_name,\
    #                           host=conf['host'],user=conf['username'],password=conf['password'])
    
    # print("all_tables")
    # print(all_tables)
    table_name_class_name=string_util.to_upper_camel_case(table_name)
    # print(f"public class Blog {")
    row_val_list=[]
    COLUMN_NAME_list=[]
    COLUMN_NAME_at_list=[]
    AddWithValue_row_list=[]
    COLUMN_NAME_update_set_list=[]
    COLUMN_COMMENT_list=[]
    COLUMN_COMMENT_show_list=[]
    if_not_null_select_list=[]
    for col_info in cols:
        COLUMN_NAME=col_info['COLUMN_NAME']
        COLUMN_COMMENT=col_info['COLUMN_COMMENT']
        
        # COLUMN_COMMENT
        COLUMN_NAME_list.append(COLUMN_NAME)
        COLUMN_NAME_at_list.append(f"@{COLUMN_NAME}")
        if COLUMN_NAME!="id":
            COLUMN_NAME_update_set_list.append(f"{COLUMN_NAME} = @{COLUMN_NAME}")
        # birthday=@birthday
        # to_lower_camel_case()
        
        if_tpl="""
        if (postData.{COLUMN_NAME} != null)
        {
            values.Add($" {COLUMN_NAME} = @{COLUMN_NAME} ");
            com.Parameters.AddWithValue("@{COLUMN_NAME}", postData.{COLUMN_NAME});
        }
        """
        if_not_null_select=if_tpl.replace("{COLUMN_NAME}",COLUMN_NAME)
        if_not_null_select_list.append(if_not_null_select)
        col_info['lower_camel_case_COLUMN_NAME']=string_util.to_lower_camel_case(COLUMN_NAME)
        col_info['upper_camel_case_COLUMN_NAME']=string_util.to_upper_camel_case(COLUMN_NAME)
        lower_camel_case_COLUMN_NAME=col_info['lower_camel_case_COLUMN_NAME']
        # input_div=f'<el-input v-model="query.{COLUMN_NAME}"></el-input>'
        # queryParam
        # query
        AddWithValue_row=f'com.Parameters.AddWithValue("@{COLUMN_NAME}", line.{COLUMN_NAME});'
        AddWithValue_row_list.append(AddWithValue_row)
        
        COLUMN_COMMENT=col_info['COLUMN_COMMENT']
        # DATA_TYPE=col_info['DATA_TYPE']
        COLUMN_COMMENT_show=COLUMN_COMMENT if COLUMN_COMMENT else COLUMN_NAME
        COLUMN_COMMENT_show_list.append(COLUMN_COMMENT_show)
        # if DATA_TYPE:
        # public class Blog {
        #     [Column(IsIdentity = true, IsPrimary = true)]
        #     public int BlogId { get; set; }
        #     public string Url { get; set; }
        #     public int Rating { get; set; }
        # }
        # DATA_TYPE
        FieldMysqlMapEntity=toFieldMysqlMapEntity(col_info)

        row_val_list.append(FieldMysqlMapEntity)
        row_val_list.append(f"line.{COLUMN_NAME} = {COLUMN_NAME};")

    field_rows="\n".join(row_val_list)
    

    COLUMN_NAMEs=", ".join(COLUMN_NAME_list)
    COLUMN_NAME_at_list_str=", ".join(COLUMN_NAME_at_list)
    AddWithValue_row_list_str="\n".join(AddWithValue_row_list)
    COLUMN_NAME_update_set_list_str=", ".join(COLUMN_NAME_update_set_list)
    if_not_null_select_list_str="\n".join(if_not_null_select_list)
    COLUMN_COMMENT_show_list_str="、 ".join(COLUMN_COMMENT_show_list)
    class_file_tpl="""
    {TABLE_COMMENT_show}编辑，可以编辑{COLUMN_COMMENT_show_list_str}等信息。
    {TABLE_COMMENT_show}列表，可以分页查看{COLUMN_COMMENT_show_list_str}等，可以进行编辑、删除等操作。
    可以通过ID对信息进行排序，对 进行模糊搜索查询
       
        """
    return class_file_tpl.replace("{table_name_class_name}",table_name_class_name).\
        replace("{field_rows}",field_rows)\
        .replace("{COLUMN_NAMEs}",COLUMN_NAMEs)\
        .replace("{COLUMN_NAME_at_list_str}",COLUMN_NAME_at_list_str)\
        .replace("{AddWithValue_row_list_str}",AddWithValue_row_list_str)\
        .replace("{table_name}",table_name)\
        .replace("{COLUMN_NAME_update_set_list_str}",COLUMN_NAME_update_set_list_str)\
        .replace("{if_not_null_select_list_str}",if_not_null_select_list_str)\
        .replace("{COLUMN_COMMENT_show_list_str}",COLUMN_COMMENT_show_list_str)\
        .replace("{TABLE_COMMENT_show}",TABLE_COMMENT_show)\
        

def select(connection,sql):
    with connection.cursor() as cursor:
        cursor.execute(sql)
        tables=cursor.fetchall()
        # print("tables")
        return tables
        # print(tables)

def get_now_time_dir(out_file_dir):
    # out_file_dir=rf"D:\毕设\lunWenDesp"
    now_time_str=time_util.get_now_time_str()
    out_file_dir=f"{out_file_dir}/{now_time_str}"
    os.makedirs(out_file_dir)
    return out_file_dir


def to_lunWenDesp():
    # do_write=False
    do_write=True
    connection = pymysql.connect(host=host,
    user=username,
    password=password,
    db=db_name,
    # db=‘数据库名称’,
    charset='utf8',
    cursorclass=pymysql.cursors.DictCursor)
    # tables=[]
    """
    {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'all_visit_count', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 31, 'AVG_ROW_LENGTH': 528, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 56), 'UPDATE_TIME': datetime.datetime(2022, 12, 13, 10, 3, 26), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'correct_count', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 10, 'AVG_ROW_LENGTH': 1638, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 57), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'student', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 48, 'AVG_ROW_LENGTH': 341, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 217, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 57), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'visit_time', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 140, 'AVG_ROW_LENGTH': 117, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 57), 'UPDATE_TIME': datetime.datetime(2022, 12, 13, 10, 3, 26), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'mqtt_control', 'TABLE_NAME': 'backdirlog', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 25, 'AVG_ROW_LENGTH': 655, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 35, 'CREATE_TIME': datetime.datetime(2022, 6, 9, 14, 57, 40), 'UPDATE_TIME': datetime.datetime(2022, 11, 19, 9, 37, 28), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'mqtt_control', 'TABLE_NAME': 'git_down', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 4, 'AVG_ROW_LENGTH': 4096, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 5, 'CREATE_TIME': datetime.datetime(2022, 6, 9, 21, 56, 42), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}]"""
    with connection.cursor() as cursor:
        cursor.execute(SELECT_table_names_sql)
        tables=cursor.fetchall()
        print("tables")
        print(tables)
        # table = cursor.fetchall() # 获取所有（输出是列表，列表中的每个元素是字典，字典的KEY是Tables_in_拼接数据库名称）
        # for i in table:
        #     print(i)
        # for k in range(len(table)): # 提取表名
        #     # tables.append(table[k]['Tables_in_数据库名称'])
        #     tables.append(table[k])
    # "private"
    print("len(tables)")
    print(len(tables))
    table_name_list=[]
    # 试卷列表，可以查看学科、等，可以进行编辑、删除等操作
    out_file_dir=rf"D:\毕设\lunWenDesp"
    # now_time_str=time_util.get_now_time_str()
    # out_file_dir=f"{out_file_dir}/{now_time_str}"
    out_file_dir=get_now_time_dir(out_file_dir)
    for table_obj in tables:
        # TABLE_NAME
        # table_name
        table_name=table_obj['TABLE_NAME']
        print(table_name)
        table_name_list.append(table_name)
        sql = f"""
        select * from information_schema.COLUMNS WHERE TABLE_SCHEMA = '{db_name}' and TABLE_NAME = '{table_name}'
        """
        cols=select(connection=connection,sql=sql)
        # cursor.execute(sql)
        # cols = cursor.fetchall()
        res=genTableLunWen(table_obj,cols)
        print(res)
        # os 
        abs_path=os.path.join(out_file_dir,f"{table_name}.md")
        print(abs_path)
        if do_write:
            with open(abs_path, "w", encoding="utf-8") as f:
                f.write(res)
    connection.close()
    print("table_name_list")
    print(table_name_list)

import os
if __name__ =="__main__":
    
    # dict list to csv 
    # table_name='product'


    # # .....
    # table_name='component'
    # print("table_name")
    # print(table_name)
    # # db_name="fastlink"
    # db_name="public"
    # # csv 中文  乱码 
    # # COLUMN_NAME
    # all_tables=get_all_tables(db_name=db_name,\
    #                           host=conf['host'],user=conf['username'],password=conf['password'])
    # print(all_tables)
    # out_file_dir=rf"D:\毕设"
    # out_file_name=f"{out_file_dir}/{db_name}.csv"
    # print("out_file_name")
    # print(out_file_name)
    # # file_util.to_csv(all_tables,out_file_name)
    # df=pd.DataFrame.from_records(all_tables)
    # # df 只要某几列 
    # # df  不同的 tablename 写入不同的 csv文件
    # # need_columns=['COLUMN_NAME','DATA_TYPE','COLUMN_COMMENT','IS_NULLABLE','COLUMN_KEY','TABLE_NAME']
    # need_columns=['字段','类型','是否为空','备注','键','COLUMN_NAME','DATA_TYPE','COLUMN_COMMENT','IS_NULLABLE','COLUMN_KEY','TABLE_NAME']
    # # df['类型']=df['DATA_TYPE']
    # # df groupby tablename 
    # # data.groupby("company")
    # # df['TABLE_NAME']
    # TABLE_NAME_set=set(df['TABLE_NAME'])
    # # df=df_chinese_name_set(df)
    # df=df_chinese_name_set(df)
    
    # for TABLE_NAME in TABLE_NAME_set:
    #     # print(TABLE_NAME)
    #     df_table=df[df['TABLE_NAME']==TABLE_NAME]
    #     df_table=df_table[need_columns]
    #     print(df_table)
    #     pandas_util.to_csv(df_table,f"{out_file_dir}/{TABLE_NAME}.csv")
    # # 。。。。。。。



    # df_groups=df.groupby("TABLE_NAME")
    # for i in df_groups:
    #     print(i)
        # i.table_name=i[0]
    
    # df['类型']=df['COLUMN_TYPE']
    # # COLUMN_TYPE
    # df['是否为空']=df['IS_NULLABLE']
    # df['键']=df['COLUMN_KEY']
    # df['备注']=df['COLUMN_COMMENT']
    # df['字段']=df['COLUMN_NAME']
    # # COLUMN_TYPE
    # # 字段
    # df=df[need_columns]
    # # excel 复制到 word 是 表格 没有表格线 
    # # 类型	是否为空	键	备注
    # print(df)
    # pan 
    
    # pandas_util.to_csv(df,out_file_name)
    # D:\codeGen\genCSharp
    # genOutDir=rf"D:\codeGen\genCSharp"
    # # res=genMysqlMapToEntity()
    # # res=genCSharpClass()
    # # res=genCreate()
    # # # genName="genCSharpClass"
    # # genName="genCreate"

    # res=genPut()
    # genName="genPut"

    # print(res)

    
    # now_time_str=time_util.get_now_time_str()

    # # import
    # # def get_now_time_str():
    # #     now_time_str = time.strftime("%Y_%m_%d_%H_%M_%S", time.localtime())
    # #     return now_time_str


    # # get_now_time_str=get_now_time()


    # outFileName=f"{genOutDir}/{table_name}_{genName}_{now_time_str}.cs"
    # print("outFileName")
    # print(outFileName)

    # with open(outFileName, "w") as f:
    #     f.write(res)

    # with open(f"{table_name}_genMysqlMapToEntity.cs", "w") as f:
    #     f.write(res)

    # res=genLunWenDesp()
    # print(res)


    # # ........ ........ . . . .. 
    # connection = pymysql.connect(host=host,
    # user=username,
    # password=password,
    # db=db_name,
    # # db=‘数据库名称’,
    # charset='utf8',
    # cursorclass=pymysql.cursors.DictCursor)
    # # tables=[]
    # """
    # {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'all_visit_count', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 31, 'AVG_ROW_LENGTH': 528, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 56), 'UPDATE_TIME': datetime.datetime(2022, 12, 13, 10, 3, 26), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'correct_count', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 10, 'AVG_ROW_LENGTH': 1638, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 57), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'student', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 48, 'AVG_ROW_LENGTH': 341, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 217, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 57), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'problemsolvingsystem', 'TABLE_NAME': 'visit_time', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 140, 'AVG_ROW_LENGTH': 117, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': None, 'CREATE_TIME': datetime.datetime(2022, 10, 30, 12, 46, 57), 'UPDATE_TIME': datetime.datetime(2022, 12, 13, 10, 3, 26), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'mqtt_control', 'TABLE_NAME': 'backdirlog', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 25, 'AVG_ROW_LENGTH': 655, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 35, 'CREATE_TIME': datetime.datetime(2022, 6, 9, 14, 57, 40), 'UPDATE_TIME': datetime.datetime(2022, 11, 19, 9, 37, 28), 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8mb4_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}, {'TABLE_CATALOG': 'def', 'TABLE_SCHEMA': 'mqtt_control', 'TABLE_NAME': 'git_down', 'TABLE_TYPE': 'BASE TABLE', 'ENGINE': 'InnoDB', 'VERSION': 10, 'ROW_FORMAT': 'Dynamic', 'TABLE_ROWS': 4, 'AVG_ROW_LENGTH': 4096, 'DATA_LENGTH': 16384, 'MAX_DATA_LENGTH': 0, 'INDEX_LENGTH': 0, 'DATA_FREE': 0, 'AUTO_INCREMENT': 5, 'CREATE_TIME': datetime.datetime(2022, 6, 9, 21, 56, 42), 'UPDATE_TIME': None, 'CHECK_TIME': None, 'TABLE_COLLATION': 'utf8_general_ci', 'CHECKSUM': None, 'CREATE_OPTIONS': '', 'TABLE_COMMENT': '', 'MAX_INDEX_LENGTH': 0, 'TEMPORARY': 'N'}]"""
    # with connection.cursor() as cursor:
    #     cursor.execute(SELECT_table_names_sql)
    #     tables=cursor.fetchall()
    #     print("tables")
    #     print(tables)
    #     # table = cursor.fetchall() # 获取所有（输出是列表，列表中的每个元素是字典，字典的KEY是Tables_in_拼接数据库名称）
    #     # for i in table:
    #     #     print(i)
    #     # for k in range(len(table)): # 提取表名
    #     #     # tables.append(table[k]['Tables_in_数据库名称'])
    #     #     tables.append(table[k])
    # # "private"
    # print("len(tables)")
    # print(len(tables))
    # table_name_list=[]
    # out_file_dir=rf"D:\毕设\lunWenDesp"
    # for table_obj in tables:
    #     # TABLE_NAME
    #     # table_name
    #     table_name=table_obj['TABLE_NAME']
    #     print(table_name)
    #     table_name_list.append(table_name)
    #     sql = f"""
    #     select * from information_schema.COLUMNS WHERE TABLE_SCHEMA = '{db_name}' and TABLE_NAME = '{table_name}'
    #     """
    #     cols=select(connection=connection,sql=sql)
    #     # cursor.execute(sql)
    #     # cols = cursor.fetchall()
    #     res=genTableLunWen(table_obj,cols)
    #     print(res)
    #     # os 
    #     abs_path=os.path.join(out_file_dir,f"{table_name}.md")
    #     print(abs_path)
    #     with open(abs_path, "w", encoding="utf-8") as f:
    #         f.write(res)
    # connection.close()
    # print("table_name_list")
    # print(table_name_list)

    # # .....  . .. . . .. . .. . .

    # mysql_table_meta_to_csv()
    to_lunWenDesp()
    # 'TABLE_NAME'