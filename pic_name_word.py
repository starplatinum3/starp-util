import docx

# 打开Word文档
input_docx_path=rf"D:\毕业设计-starp-考试.docx"
# doc = docx.Document(rf'D:\毕设\毕业设计-starp-考试.doc')
doc = docx.Document(input_docx_path)


from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement

def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para
# 图片计数器
pic_count = 0
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def pic_center_set_up(run,para):
    # AttributeError: 'Run' object has no attribute 'inline_shapes'
    # if not inline_shapes in run:
    #     return
    # TypeError: xpath() got an unexpected keyword argument 'namespaces'
    if not run._element.xpath('.//wp:docPr'):
        return
    # if not run._element.xpath('.//wp:docPr', namespaces=run._element.nsmap):
    #     return
    if not run.inline_shapes:
        return
    # 获取图片对象
    inline_shape = run.inline_shapes[0]
    # 设置图片对象的对齐方式为居中对齐
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 获取图片的xml代码
    xml = inline_shape._inline.graphic.graphicData.pic._pic.getroot().xml
    # 设置图片水平方向的位置为居中
    xml = xml.replace('<wp:posOffset', '<wp:positionH relativeFrom="margin"><wp:posOffset')
    # 设置图片垂直方向的位置为居中
    xml = xml.replace('<wp:positionH', '<wp:positionV relativeFrom="margin"><wp:posOffset')
    # 重新设置图片的xml代码
    inline_shape._inline.graphic.graphicData.pic._pic.clear()
    inline_shape._inline.graphic.graphicData.pic._pic._element.xml = xml

# AttributeError: 'Paragraph' object has no attribute 'insert_paragraph_after'
# 遍历文档中的每个段落和图片
for para in doc.paragraphs:
    # 判断段落中是否有图片
    for run in para.runs:
        if run._element.tag.endswith('}r'):
            for elem in run._element:
                if elem.tag.endswith('}pict'):
                    # # 获取图片对象
                    # inline_shape = run.inline_shapes[0]
                    # # 设置图片对象的对齐方式为居中对齐
                    # para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # # 获取图片的xml代码
                    # xml = inline_shape._inline.graphic.graphicData.pic._pic.getroot().xml
                    # # 设置图片水平方向的位置为居中
                    # xml = xml.replace('<wp:posOffset', '<wp:positionH relativeFrom="margin"><wp:posOffset')
                    # # 设置图片垂直方向的位置为居中
                    # xml = xml.replace('<wp:positionH', '<wp:positionV relativeFrom="margin"><wp:posOffset')
                    # # 重新设置图片的xml代码
                    # inline_shape._inline.graphic.graphicData.pic._pic.clear()
                    # inline_shape._inline.graphic.graphicData.pic._pic._element.xml = xml
                    pic_center_set_up(run=run,para=para)
                    pic_count += 1
                    # 在段落下面插入一个新段落
                    # new_para = para.insert_paragraph_after('')
                    # 图片 居中
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
                    new_para =insert_paragraph_after(para,'')
                    # 在新段落中插入自动编号的标签
                    # new_para.add_run 居中
                    pic_caption_run=new_para.add_run('图{}.{}'.format(pic_count // 10 + 1, pic_count % 10 + 1))
                    # pic_caption_run.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
                    # pic_caption_run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
                    new_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
                    # 将图片插入到新段落中
                    # new_run = new_para.add_run()
                    # new_run.add_picture('example{}.jpg'.format(pic_count))

# 保存文档
# doc.save('example_with_caption.docx')\
# 
import time_util
now_time_str=time_util.get_now_time_str()

save_path=rf'D:\毕设\out\example_with_captions_{now_time_str}.docx'
print("save_path")
print(save_path)
doc.save(save_path)
import os 

import time 

# time.sleep(4)

# os.system(f"explorer {save_path}")
os.system(f"start {save_path}")