
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement

def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para



from docx import Document
import docx

def run_add_fldChar(run,cmd_str):
    fldChar = docx.oxml.OxmlElement('w:fldChar')
    fldChar.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText_STYLEREF = docx.oxml.OxmlElement('w:instrText')
    # instrText.text = f' SEQ {target} \\* ARABIC'
    # instrText_STYLEREF.text = fr' STYLEREF 1\s '
    instrText_STYLEREF.text = cmd_str
    # 表 0.2
    run._r.append(instrText_STYLEREF)
    # run._r.append('.')
    # run.

    # instrTextSeq = docx.oxml.OxmlElement('w:instrText')
    # # instrText.text = f' SEQ {target} \\* ARABIC'
    # instrTextSeq.text = fr' SEQ {target} \* ARABIC \s 1'
    # 表 0.2
    # run._r.append(instrTextSeq)

    fldChar = docx.oxml.OxmlElement('w:fldChar')
    fldChar.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

#  table: docx.table.Table,
def add_caption(doc, caption: str,target = 'Table'):
    """
    Based on: https://github.com/python-openxml/python-docx/issues/359
    """
    # target = 'Table'
    # doc
    # caption type
    paragraph = doc.add_paragraph(f'{target} ', style='Caption')

    # numbering field
    run = paragraph.add_run()
    run_add_fldChar(run,cmd_str=fr' STYLEREF 1\s ')

    fldChar = docx.oxml.OxmlElement('w:fldChar')
    fldChar.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText_STYLEREF = docx.oxml.OxmlElement('w:instrText')
    # instrText.text = f' SEQ {target} \\* ARABIC'
    instrText_STYLEREF.text = fr' STYLEREF 1\s '
    # 表 0.2
    run._r.append(instrText_STYLEREF)
    # run._r.append('.')
    # run.

    instrTextSeq = docx.oxml.OxmlElement('w:instrText')
    # instrText.text = f' SEQ {target} \\* ARABIC'
    instrTextSeq.text = fr' SEQ {target} \* ARABIC \s 1'
    # 表 0.2
    run._r.append(instrTextSeq)

    fldChar = docx.oxml.OxmlElement('w:fldChar')
    fldChar.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

    # caption text
    paragraph.add_run(f' {caption}')